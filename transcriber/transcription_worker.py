import threading
import queue
import os
import sys
import warnings
import contextlib
from io import StringIO
from typing import List, Optional, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

from .ffmpeg_utils import convert_to_wav_ffmpeg, chunk_wav_file
from .whisper_utils import load_whisper_model, transcribe_audio_segment, format_timecode
from .audio_enhancement import enhance_audio
from .diarization import get_diarization_pipeline, perform_diarization, assign_speakers_to_segments, is_diarization_available, suppress_stdout_stderr

class TranscriptionWorker(threading.Thread):
    """Process multiple files in a background thread."""
    
    def __init__(
        self,
        input_files: List[str],
        out_dir: str,
        model_name: str = "medium",
        device: str = "cpu",
        sample_rate: int = 16000,
        channels: int = 1,
        chunk: bool = True,
        chunk_length: int = 300,
        log_queue: Optional[queue.Queue] = None,
        progress_queue: Optional[queue.Queue] = None,
        enhance_audio: bool = False,
        enable_diarization: bool = False,
        huggingface_token: Optional[str] = None,
        min_speakers: Optional[int] = 1,
        max_speakers: Optional[int] = 2,
        segmentation: float = 1.0,
        apply_diarization_preprocessing: bool = True
    ):
        super().__init__()
        # Set environment variables to suppress warnings
        os.environ["TOKENIZERS_PARALLELISM"] = "false"
        
        self.input_files = input_files
        self.out_dir = out_dir
        self.model_name = model_name
        self.device = device
        self.sample_rate = sample_rate
        self.channels = channels
        self.chunk = chunk
        self.chunk_length = chunk_length
        self.log_queue = log_queue
        self.progress_queue = progress_queue
        self.stop_event = threading.Event()
        self.doc_messages = []
        
        self.enhance_audio_flag = enhance_audio
        self.enable_diarization = enable_diarization
        self.huggingface_token = huggingface_token
        self.min_speakers = min_speakers
        self.max_speakers = max_speakers
        self.segmentation = segmentation
        self.apply_diarization_preprocessing = apply_diarization_preprocessing
        self.diarization_pipeline = None
        
        # Define speaker colors for document formatting
        self.speaker_colors = {
            "SPEAKER_00": RGBColor(0, 0, 139),  # Dark blue
            "SPEAKER_01": RGBColor(139, 0, 0),  # Dark red
            "SPEAKER_02": RGBColor(0, 100, 0),  # Dark green
            "SPEAKER_03": RGBColor(148, 0, 211), # Dark violet
            "SPEAKER_04": RGBColor(210, 105, 30), # Chocolate
            "SPEAKER_05": RGBColor(0, 139, 139), # Dark cyan
            "SPEAKER_06": RGBColor(169, 169, 169), # Dark gray
            "SPEAKER_07": RGBColor(184, 134, 11), # Dark goldenrod
            "SPEAKER_08": RGBColor(139, 69, 19), # Saddle brown
            "SPEAKER_09": RGBColor(85, 107, 47), # Dark olive green
            "UNKNOWN": RGBColor(128, 128, 128)   # Gray
        }

    def _log(self, msg: str) -> None:
        """Helper to push log messages to queue or print."""
        if self.log_queue:
            self.log_queue.put(msg)
        else:
            print(msg)
        self.doc_messages.append(msg)

    def _create_document(self, filename: str) -> Document:
        """Create and initialize a new document with basic styling."""
        doc = Document()
        # Add title
        title = doc.add_heading(f"Transcript: {os.path.basename(filename)}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add processing details
        details = doc.add_paragraph()
        details.add_run("Processing Details\n").bold = True
        details.add_run(f"Model: {self.model_name}\n")
        details.add_run(f"Device: {self.device}\n")
        details.add_run(f"Sample Rate: {self.sample_rate} Hz\n")
        details.add_run(f"Channels: {self.channels}\n")
        if self.chunk:
            details.add_run(f"Chunk Length: {self.chunk_length} seconds\n")
        if self.enhance_audio_flag:
            details.add_run("Audio Enhancement: Enabled\n")
        if self.enable_diarization:
            details.add_run("Speaker Diarization: Enabled\n")
            if self.min_speakers is not None and self.max_speakers is not None:
                details.add_run(f"Speaker Constraints: Min={self.min_speakers}, Max={self.max_speakers}\n")
            elif self.max_speakers is not None:
                details.add_run(f"Speaker Constraint: Max={self.max_speakers}\n")
            if self.apply_diarization_preprocessing:
                details.add_run("Diarization Preprocessing: Enabled\n")
        
        # Add a separator
        doc.add_paragraph("=" * 80)
        
        # Add transcription heading
        doc.add_heading("Transcription", level=2)
        return doc

    def _add_speaker_summary(self, doc: Document, diarization_result: Dict):
        """Add a summary of speakers at the beginning of the document."""
        if not diarization_result:
            return
            
        speaker_mapping = diarization_result.get("speaker_mapping", {})
        speakers = diarization_result.get("speakers", {})
        
        if not speakers:
            return
            
        # Add speaker summary heading
        doc.add_heading("Speaker Summary", level=2)
        
        # Create a table for the speakers
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Speaker"
        header_cells[1].text = "Speaking Time"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Calculate total speaking time for each speaker
        speaker_durations = {}
        for speaker_id, segments in speakers.items():
            total_duration = sum(segment["end"] - segment["start"] for segment in segments)
            speaker_durations[speaker_id] = total_duration
        
        # Add a row for each speaker
        for speaker_id, duration in speaker_durations.items():
            # Add a row for this speaker
            row = table.add_row().cells
            
            # Format the cell with the appropriate color
            speaker_color = self.speaker_colors.get(speaker_id, RGBColor(0, 0, 0))
            speaker_run = row[0].paragraphs[0].add_run(speaker_id)
            speaker_run.bold = True
            speaker_run.font.color.rgb = speaker_color
            
            # Add speaking time
            row[1].text = f"{duration:.1f} seconds"
        
        # Add a paragraph after the table
        doc.add_paragraph()
        
        # Add a separator
        doc.add_paragraph("=" * 80)
        
        # Add transcription heading
        doc.add_heading("Transcription", level=2)

    def _add_segment_to_doc(self, doc: Document, segment: dict, chunk_idx: Optional[int] = None, always_show_speaker: bool = True):
        """Add a single transcription segment to the document with enhanced formatting."""
        # Create paragraph for this segment
        p = doc.add_paragraph()
        
        # Add chunk indicator if provided
        if chunk_idx is not None:
            chunk_run = p.add_run(f"[Chunk {chunk_idx}] ")
            chunk_run.bold = True
            chunk_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add timestamp
        start_str = format_timecode(segment["start"])
        end_str = format_timecode(segment["end"])
        time_run = p.add_run(f"[{start_str} - {end_str}] ")
        time_run.bold = True
        time_run.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Add speaker label if available with improved formatting
        if "speaker" in segment and always_show_speaker:
            speaker = segment["speaker"]
            speaker_color = self.speaker_colors.get(speaker, RGBColor(0, 0, 0))
            
            # Just use the speaker ID
            speaker_label = f"{speaker}: "
                
            speaker_run = p.add_run(speaker_label)
            speaker_run.bold = True
            speaker_run.font.color.rgb = speaker_color
            
            # Indent the paragraph to make speaker changes more visible
            p.paragraph_format.left_indent = Inches(0.25)

        # Add transcribed text
        text_run = p.add_run(segment["text"].strip())
        text_run.font.size = Pt(11)

    def _format_transcript_with_speaker_changes(self, doc: Document, segments: List[Dict]):
        """
        Format transcript with clear speaker changes and improved readability.
        Always shows the speaker label for each segment for better readability.
        """
        last_speaker = None
        current_chunk = None
        
        for segment in segments:
            current_speaker = segment.get("speaker", "UNKNOWN")
            segment_chunk = segment.get("chunk_idx")
            
            # Add extra spacing when speaker changes or chunk changes
            if current_speaker != last_speaker or segment_chunk != current_chunk:
                # Add extra spacing for speaker change or chunk change
                if last_speaker is not None:
                    doc.add_paragraph()  # Add blank line between changes
                
                # Update tracking variables
                last_speaker = current_speaker
                current_chunk = segment_chunk
            
            # Add segment with enhanced formatting, always showing the speaker
            self._add_segment_to_doc(doc, segment, segment_chunk, always_show_speaker=True)

    def run(self) -> None:
        """Main processing loop."""
        # Suppress warnings during processing
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            
            total = len(self.input_files)
            if total == 0:
                self._log("No files to process.")
                return

            # Load Whisper model once
            model = load_whisper_model(
                self.model_name,
                self.device,
                log_queue=self.log_queue
            )
            if not model:
                self._log(f"Failed to load Whisper model '{self.model_name}'")
                return

            # Check diarization availability
            diarization_available, diarization_error = is_diarization_available()
            if self.enable_diarization and not diarization_available:
                self._log(f"Speaker diarization is not available: {diarization_error}")
                self._log("Continuing without diarization.")
                self.enable_diarization = False
                
            # Load diarization pipeline if enabled and available
            if self.enable_diarization:
                try:
                    self._log("Loading diarization pipeline...")
                    with suppress_stdout_stderr():
                        self.diarization_pipeline = get_diarization_pipeline(
                            self.huggingface_token,
                            device=self.device
                        )
                    self._log(f"Diarization pipeline loaded successfully on device: {self.device}")
                except Exception as e:
                    self._log(f"Failed to load diarization pipeline: {e}")
                    self._log("Continuing without diarization.")
                    self.enable_diarization = False

            for idx, f in enumerate(self.input_files, start=1):
                if self.stop_event.is_set():
                    self._log("Stop requested. Halting worker.")
                    break

                self._log(f"Processing file {idx}/{total}: '{f}'")
                
                # Create new document for this file
                doc = self._create_document(f)
                
                # 1) Convert to WAV
                wav_file = convert_to_wav_ffmpeg(
                    f,
                    self.sample_rate,
                    self.channels,
                    log_queue=self.log_queue
                )
                if not wav_file:
                    self._log(f"WAV conversion failed for '{f}'. Skipping.")
                    continue

                # 2) Perform diarization on the original file
                diarization_result = None
                if self.enable_diarization and self.diarization_pipeline:
                    self._log("Performing speaker diarization...")
                    with suppress_stdout_stderr():
                        diarization_result = perform_diarization(
                            wav_file,  # Use original WAV file
                            self.diarization_pipeline,
                            log_queue=self.log_queue,
                            min_speakers=self.min_speakers,
                            max_speakers=self.max_speakers,
                            segmentation=self.segmentation,
                            apply_preprocessing=self.apply_diarization_preprocessing
                        )
                    num_speakers = diarization_result.get("num_speakers", 0)
                    self._log(f"Diarization complete. Detected {num_speakers} speakers.")
                    
                    # Add speaker summary to document
                    self._add_speaker_summary(doc, diarization_result)

                # 3) Split into chunks if needed
                if self.chunk:
                    chunks = chunk_wav_file(
                        wav_file,
                        self.chunk_length,
                        log_queue=self.log_queue
                    )
                else:
                    chunks = [wav_file]

                # 4) Process all chunks
                all_segments = []
                for chunk_idx, chunk_path in enumerate(chunks, start=1):
                    if self.stop_event.is_set():
                        self._log("Stop requested during chunk processing.")
                        break
                    
                    # Enhance audio for transcription if enabled
                    if self.enhance_audio_flag:
                        self._log(f"Enhancing chunk {chunk_idx}/{len(chunks)} for transcription...")
                        try:
                            enhanced_path = enhance_audio(chunk_path)
                            audio_to_transcribe = enhanced_path
                        except Exception as e:
                            audio_to_transcribe = chunk_path
                            self._log(f"Chunk enhancement failed: {e}. Using original audio.")
                    else:
                        audio_to_transcribe = chunk_path

                    self._log(f"Transcribing chunk {chunk_idx}/{len(chunks)}...")
                    result = transcribe_audio_segment(
                        model,
                        audio_to_transcribe,
                        verbose=True,
                        log_queue=self.log_queue
                    )
                    
                    # Apply speaker diarization to segments if enabled
                    if self.enable_diarization and diarization_result:
                        self._log(f"Assigning speakers to segments in chunk {chunk_idx}...")
                        result["segments"] = assign_speakers_to_segments(
                            result["segments"],
                            diarization_result
                        )
                    
                    # Add chunk index to segments for tracking
                    for segment in result["segments"]:
                        segment["chunk_idx"] = chunk_idx if len(chunks) > 1 else None
                    
                    # Add segments to all_segments for overall processing
                    all_segments.extend(result["segments"])

                # Format the transcript with improved speaker clarity
                self._format_transcript_with_speaker_changes(doc, all_segments)

                # Add processing log if we have messages
                if self.doc_messages:
                    doc.add_page_break()
                    doc.add_heading("Processing Log", level=2)
                    log_para = doc.add_paragraph()
                    for msg in self.doc_messages:
                        log_para.add_run(msg + "\n")

                # Save document
                base = os.path.splitext(os.path.basename(f))[0]
                out_path = os.path.join(self.out_dir, base + ".docx")
                doc.save(out_path)
                self._log(f"Saved transcript to '{out_path}'")

                # Clear messages for next file
                self.doc_messages = []

                # Update progress
                if self.progress_queue:
                    self.progress_queue.put((idx, total))

            self._log("Transcription worker finished.")

    def stop(self) -> None:
        """Signal the thread to stop gracefully."""
        self.stop_event.set()
