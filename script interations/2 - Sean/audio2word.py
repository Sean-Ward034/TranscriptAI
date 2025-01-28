import os
import subprocess
from docx import Document
# from pydub import AudioSegment
# import moviepy.editor as mp
import whisper

def convert_to_wav_ffmpeg(input_file):
    """
    Converts any audio file to WAV format using FFmpeg (no pydub).
    """
    output_file = os.path.splitext(input_file)[0] + ".wav"
    if os.path.isfile(output_file):
        os.remove(output_file)

    # FFmpeg command: extract audio to WAV
    command = [
        "ffmpeg", "-y",
        "-i", input_file,
        "-vn",            # no video
        "-acodec", "pcm_s16le",
        "-ar", "16000",   # sample rate
        "-ac", "1",       # mono
        output_file
    ]

    try:
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"Successfully converted {input_file} to {output_file}")
    except subprocess.CalledProcessError as e:
        print(f"Error converting {input_file} to WAV: {e}")

def convert_video_to_wav_ffmpeg(video_file):
    """
    Converts a video file to WAV audio using FFmpeg (no moviepy).
    """
    output_file = os.path.splitext(video_file)[0] + ".wav"
    if os.path.isfile(output_file):
        os.remove(output_file)

    # FFmpeg command: extract audio track to WAV
    command = [
        "ffmpeg", "-y",
        "-i", video_file,
        "-vn",            # ignore video
        "-acodec", "pcm_s16le",
        "-ar", "16000",   # sample rate
        "-ac", "1",       # mono
        output_file
    ]

    try:
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print(f"Successfully converted video {video_file} to {output_file}")
    except subprocess.CalledProcessError as e:
        print(f"Error converting video {video_file}: {e}")

'''
# OLD pydub-based convert_to_wav function, now replaced by ffmpeg
def convert_to_wav(input_file):
    # Converts any audio file to WAV format.
    try:
        audio = AudioSegment.from_file(input_file)
        output_file = input_file.rsplit('.', 1)[0] + ".wav"
        audio.export(output_file, format="wav")
        print(f"Successfully converted {input_file} to {output_file}")
    except Exception as e:
        print(f"Error converting {input_file}: {e}")
'''

'''
# OLD moviepy-based convert_video_to_wav function, replaced by ffmpeg
def convert_video_to_wav(video_path, output_path):
    # Converts a video file to WAV audio.
    video_clip = mp.VideoFileClip(video_path)
    audio_clip = video_clip.audio
    audio_clip.write_audiofile(output_path, codec='pcm_s16le')
'''

def transcribe_audio(audio_file):
    """
    Transcribes an audio file using the Whisper large model
    (instead of base, for potentially better accuracy).
    """
    # model = whisper.load_model("base")  # OLD
    model = whisper.load_model("medium")  # NEW
    result = model.transcribe(audio_file)
    return result["text"]

'''
# OLD process_directory function – replaced below with an updated version
def process_directory(directory):
    for filename in os.listdir(directory):
        if filename.endswith(".wav") or filename.endswith(".mp3") or filename.endswith(".m4a"):
            ...
            # pydub-based conversion
            # ...
            # transcribe_audio
            # ...
'''

def process_directory(
    directory,
    output_directory,
    file_types=("audio", "video"),
    single_doc=False
):
    """
    Transcribes all audio or video files in a directory and saves
    the results to Word documents, using only FFmpeg conversions.
    
    If single_doc=True, all transcriptions go into one docx;
    otherwise each file gets its own docx.
    """

    combined_doc = Document() if single_doc else None

    audio_ext = {".mp3", ".m4a", ".wav", ".aac", ".ogg"}
    video_ext = {".mp4", ".avi", ".mkv", ".mov", ".flv", ".wmv"}

    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        base, ext = os.path.splitext(filename.lower())

        if not os.path.isfile(file_path):
            continue

        # Decide if it's audio or video
        if ext in audio_ext and "audio" in file_types:
            if ext != ".wav":
                convert_to_wav_ffmpeg(file_path)
                file_path = os.path.splitext(file_path)[0] + ".wav"

        elif ext in video_ext and "video" in file_types:
            convert_video_to_wav_ffmpeg(file_path)
            file_path = os.path.splitext(file_path)[0] + ".wav"
        else:
            # Skip files that don't match audio or video
            continue

        # Transcribe using Whisper
        transcript = transcribe_audio(file_path)
        print(f"\n=== Transcript for {filename} ===\n{transcript}")

        # If using a single doc
        if single_doc and combined_doc is not None:
            combined_doc.add_heading(filename, level=2)
            combined_doc.add_paragraph(transcript)
        else:
            # Create individual doc
            doc = Document()
            doc.add_paragraph(transcript)
            doc_name = os.path.join(output_directory, base + ".docx")
            doc.save(doc_name)

    # If single_doc is True, save everything in one doc at the end
    if single_doc and combined_doc is not None:
        combined_path = os.path.join(output_directory, "combined_transcripts.docx")
        combined_doc.save(combined_path)

if __name__ == "__main__":
    # HARDCODED path from older scripts (commented out)
    # directory_to_process = "C:\\Backups\\Files\\Wav\\"

    # Prompt the user for input directory
    directory_to_process = input("Enter the directory to scan for audio/video: ").strip()
    while not os.path.isdir(directory_to_process):
        directory_to_process = input("Please enter a valid directory path: ").strip()

    # Prompt for output directory
    output_directory = input("Enter the directory to save transcribed Word docs: ").strip()
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    # Ask user if they want to process audio, video, or both
    choice = input("Process [A]udio, [V]ideo, or [B]oth? (Default B) ").strip().lower()
    if choice == 'a':
        file_types = ("audio",)
    elif choice == 'v':
        file_types = ("video",)
    else:
        file_types = ("audio", "video")

    # Ask user if they want a single combined document
    single_choice = input("Combine all transcriptions in one .docx file? (Y/N, Default=N) ").strip().lower()
    single_doc = (single_choice == 'y')

    # Process the directory
    process_directory(
        directory=directory_to_process,
        output_directory=output_directory,
        file_types=file_types,
        single_doc=single_doc
    )
